import markdown
from docx import Document
from bs4 import BeautifulSoup
import os

def markdown_to_docx(input_file, output_file):
    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        html_content = markdown.markdown(md_content, extensions=['extra', 'tables'])

        doc = Document()

        soup = BeautifulSoup(html_content, 'html.parser')

        for element in soup:
            if not hasattr(element, 'name') or not element.name:
                continue

            if element.name == 'h1':
                doc.add_heading(element.text.strip(), level=1)
            elif element.name == 'h2':
                doc.add_heading(element.text.strip(), level=2)
            elif element.name == 'h3':
                doc.add_heading(element.text.strip(), level=3)
            elif element.name == 'h4':
                doc.add_heading(element.text.strip(), level=4)
            elif element.name == 'h5':
                doc.add_heading(element.text.strip(), level=5)
            elif element.name == 'h6':
                doc.add_heading(element.text.strip(), level=6)
            elif element.name == 'p':
                stripped_text = element.text.strip()
                if stripped_text:
                     doc.add_paragraph(stripped_text)
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.text.strip(), style='List Bullet')
            elif element.name == 'ol':
                for li in element.find_all('li', recursive=False):
                     doc.add_paragraph(li.text.strip(), style='List Number')
            elif element.name == 'table':
                rows = element.find_all('tr')
                if not rows:
                    continue

                first_row_cells = rows[0].find_all(['th', 'td'])
                if not first_row_cells:
                    continue
                num_cols = len(first_row_cells)

                table = doc.add_table(rows=0, cols=num_cols)
                table.style = 'Table Grid'

                for row_element in rows:
                    cells_data = row_element.find_all(['th', 'td'])
                    docx_row_cells = table.add_row().cells
                    for j, cell_element in enumerate(cells_data):
                         if j < num_cols:
                            docx_row_cells[j].text = cell_element.text.strip()


        doc.save(output_file)
        print(f"Файл успешно сохранён как {output_file}")

    except FileNotFoundError:
        print(f"Error: File '{input_file}' not found.")
    except Exception as e:
        print(f"Error: {str(e)}")

if __name__ == "__main__":
    input_md = input("Введите имя входного Markdown файла (например: document.md): ")

    if not input_md.lower().endswith('.md'):
        print("Предупреждение: имя входного файла не заканчивается на .md")

    base_name = os.path.splitext(input_md)[0]
    output_docx = f"{base_name}.docx"

    markdown_to_docx(input_md, output_docx)
